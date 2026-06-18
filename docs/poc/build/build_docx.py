"""Build the printable, professional Word (.docx) version of the Atlas POC.

Run after render_diagrams.py. Produces:
    docs/poc/Atlas-Unified-Knowledge-Platform-POC.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "..", "assets")
OUT = os.path.join(HERE, "..", "Atlas-Unified-Knowledge-Platform-POC.docx")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLUE = RGBColor(0x2E, 0x5E, 0x8C)
GOLD = RGBColor(0xB7, 0x8A, 0x2E)
GREY = RGBColor(0x5B, 0x64, 0x70)
DARK = RGBColor(0x22, 0x22, 0x22)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HDR_FILL = "1F3A5F"     # navy table header
ALT_FILL = "EAF1F8"     # light row stripe
BODY_FONT = "Calibri"
HEAD_FONT = "Calibri Light"

FIG = lambda n: os.path.join(ASSETS, n)

doc = Document()

# ---------------------------------------------------------------------------
# Base styles
# ---------------------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
pf = normal.paragraph_format
pf.space_after = Pt(8)
pf.line_spacing = 1.12


def _style_heading(name, size, color, bold=True, space_before=14, space_after=6):
    st = doc.styles[name]
    st.font.name = HEAD_FONT
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(space_before)
    st.paragraph_format.space_after = Pt(space_after)
    st.paragraph_format.keep_with_next = True


_style_heading("Title", 30, NAVY)
_style_heading("Heading 1", 17, NAVY)
_style_heading("Heading 2", 13.5, BLUE)
_style_heading("Heading 3", 11.5, BLUE)


# ---------------------------------------------------------------------------
# Page setup: margins + header/footer with page numbers
# ---------------------------------------------------------------------------
def set_margins(section, top=1.0, bottom=1.0, left=1.0, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def _field(paragraph, instr):
    run = paragraph.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    tx = OxmlElement("w:t"); tx.text = "1"
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    run._r.append(fb); run._r.append(it); run._r.append(fs); run._r.append(tx); run._r.append(fe)
    return run


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ts in (Inches(3.25), Inches(6.5)):
        p.paragraph_format.tab_stops.add_tab_stop(ts, WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run("Atlas — Unified Knowledge Platform  ·  Technical & Architectural POC\t")
    r.font.size = Pt(8); r.font.color.rgb = GREY; r.font.name = BODY_FONT
    rp = p.add_run("Page ")
    rp.font.size = Pt(8); rp.font.color.rgb = GREY
    _field(p, "PAGE").font.size = Pt(8)
    rt = p.add_run(" of ")
    rt.font.size = Pt(8); rt.font.color.rgb = GREY
    _field(p, "NUMPAGES").font.size = Pt(8)


def add_header(section, text):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.size = Pt(8); r.font.color.rgb = GREY; r.italic = True; r.font.name = BODY_FONT


# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------
def h1(text):
    doc.add_paragraph(text, style="Heading 1")


def h2(text):
    doc.add_paragraph(text, style="Heading 2")


def h3(text):
    doc.add_paragraph(text, style="Heading 3")


def para(text, italic=False, size=11, color=DARK, align=None, space_after=8, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold
    r.font.size = Pt(size); r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(parts, space_after=8):
    """parts: list of (text, bold)"""
    p = doc.add_paragraph()
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, lead=None, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if lead:
        r = p.add_run(lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def olist(items):
    """Ordered list with explicit, restarting numbers and a hanging indent.

    items: list of either "text" or ("lead", "text").
    """
    for i, item in enumerate(items, start=1):
        lead, text = (item if isinstance(item, tuple) else (None, item))
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(3)
        rn = p.add_run(f"{i}.\t"); rn.bold = True
        if lead:
            rl = p.add_run(lead); rl.bold = True
        p.add_run(text)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, color=DARK, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color
        r.font.name = BODY_FONT


def table(headers, rows, widths=None, font_size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        _set_cell_text(hdr[i], htext, bold=True, color=WHITE, size=font_size)
        _shade(hdr[i], HDR_FILL)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            _set_cell_text(cells[i], val, size=font_size)
            if ridx % 2 == 1:
                _shade(cells[i], ALT_FILL)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


_fig_counter = {"n": 0}


def figure(filename, caption, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(FIG(filename), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(10)


def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "1F3A5F")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def page_break():
    doc.add_page_break()


def callout(text):
    """A shaded, bordered emphasis box (single-cell table)."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    _shade(cell, ALT_FILL)
    cell.width = Inches(6.5)
    _set_cell_text(cell, text, size=10.5, color=NAVY)
    cell.paragraphs[0].runs[0].italic = True
    # left accent border
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "0"); left.set(qn("w:color"), "C8A24B")
    borders.append(left); tcPr.append(borders)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ===========================================================================
# COVER PAGE
# ===========================================================================
sec = doc.sections[0]
set_margins(sec)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ATLAS"); r.font.name = HEAD_FONT; r.font.size = Pt(46)
r.font.bold = True; r.font.color.rgb = NAVY

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Unified Knowledge Platform"); r.font.name = HEAD_FONT
r.font.size = Pt(20); r.font.color.rgb = BLUE

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Ask once. Receive one clear, consolidated, cited answer —")
r.font.size = Pt(12); r.font.color.rgb = GREY; r.italic = True
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("from every relevant page, across every domain and project.")
r.font.size = Pt(12); r.font.color.rgb = GREY; r.italic = True

doc.add_paragraph()
hrule()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Technical & Architectural Proof of Concept")
r.font.name = HEAD_FONT; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = DARK
hrule()

for _ in range(2):
    doc.add_paragraph()

meta = doc.add_table(rows=0, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in [
    ("Prepared for", "Client Executive & Architecture Review"),
    ("Prepared by", "CJ-Covance"),
    ("Document type", "Proof of Concept — Proposal"),
    ("Version", "1.0"),
    ("Date", "June 2026"),
    ("Classification", "Confidential"),
]:
    row = meta.add_row().cells
    _set_cell_text(row[0], k, bold=True, color=NAVY, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_text(row[1], v, size=11)
    row[0].width = Inches(2.3); row[1].width = Inches(3.3)

page_break()

# ===========================================================================
# DOCUMENT CONTROL + TABLE OF CONTENTS
# ===========================================================================
h1("Document Control")
table(
    ["Version", "Date", "Author", "Status", "Notes"],
    [["1.0", "June 2026", "CJ-Covance", "Draft for client review",
      "Initial technical & architectural POC proposal."]],
    widths=[0.9, 1.1, 1.6, 1.7, 2.0],
)
para("This document is confidential and intended solely for the named client's "
     "evaluation of the proposed proof of concept.", italic=True, size=9.5, color=GREY)

doc.add_paragraph()
h1("Contents")
toc_p = doc.add_paragraph()
_field(toc_p, 'TOC \\o "1-2" \\h \\z \\u')
para("If the contents above appear empty, right-click and choose "
     "\u201cUpdate Field\u201d (or press Ctrl+A then F9) to populate page numbers.",
     italic=True, size=9, color=GREY)

page_break()

# ===========================================================================
# 1. EXECUTIVE SUMMARY
# ===========================================================================
h1("1.  Executive Summary")

h2("1.1  The problem")
para("Across the organization, knowledge is real but fragmented. The answer a "
     "person needs is typically spread across a wiki space, a SharePoint site, "
     "several PDFs on a shared drive, a closed ticket, and an email thread — each "
     "owned by a different team, each using different vocabulary, and each "
     "requiring the right access. The cost is not storing documents; it is the "
     "human effort to find, read, reconcile, and summarize them.")
para("A single cross-team question — for example, \u201cWhat is our current "
     "data-retention policy for clinical documents, and which projects already "
     "implement it?\u201d — can take an expert hours to answer, and the result is "
     "only as good as the documents that person happened to find. The recurring "
     "symptoms are familiar:")
bullet("\u201cI know the answer exists somewhere — I just can\u2019t find it.\u201d")
bullet("\u201cEvery team documents the same thing differently.\u201d")
bullet("\u201cBy the time I\u2019ve read everything, I\u2019m unsure which version is current.\u201d")
bullet("\u201cNew joiners take months to learn where knowledge lives.\u201d")

h2("1.2  The proposed solution")
para("Atlas is a scalable, extensible enterprise platform that centralizes "
     "knowledge across domains and projects and lets people ask a question once "
     "and receive a single, clear, consolidated answer — summarized from every "
     "relevant page, with inline citations to the original sources. Atlas does "
     "three things well:")
olist([
    ("Unify.  ", "Connectors continuously ingest content from many systems into "
     "one normalized, searchable knowledge layer."),
    ("Understand.  ", "Content is parsed, chunked, embedded, and indexed for both "
     "semantic and keyword retrieval, enriched with metadata (domain, project, "
     "owner, freshness, access)."),
    ("Answer.  ", "A retrieval-augmented generation (RAG) pipeline finds the most "
     "relevant passages across all sources and synthesizes a concise, grounded "
     "answer with citations — never a bare list of links."),
])
para("Crucially, Atlas respects existing permissions: a user only ever sees "
     "answers built from content they are already entitled to read.")

h2("1.3  What changes for the user")
table(
    ["Today", "With Atlas"],
    [
        ["Search returns dozens of links; you open and read each one.",
         "You receive one consolidated answer, with citations to expand."],
        ["You manually reconcile conflicting documents.",
         "Atlas surfaces and flags conflicting or duplicate sources."],
        ["Knowledge lives in silos per team and per tool.",
         "One question spans every domain and project you can access."],
        ["Answers depend on who you ask.",
         "Answers are consistent, grounded, and fully traceable."],
    ],
    widths=[3.25, 3.25],
)

h2("1.4  Why this is feasible now")
bullet("Grounding answers in retrieved source text makes them auditable and "
       "citable, directly addressing enterprise trust and compliance concerns.",
       lead="RAG is enterprise-ready.  ")
bullet("New sources and domains are added as configuration plus a small "
       "connector — not a re-architecture.", lead="Connector-first design.  ")
bullet("The LLM and embedding model are swappable components, so the platform is "
       "never locked to one vendor and can run in a private cloud or on-premises.",
       lead="Model-agnostic.  ")

h2("1.5  Value and outcomes")
bullet("Time-to-answer drops from hours of manual reading to seconds.")
bullet("Every answer is grounded and cited, reducing reliance on tribal knowledge.")
bullet("Faster onboarding and reduced expert interrupt-load.")
bullet("A reusable knowledge asset that powers search, Q&A, and future automation.")

h2("1.6  The ask")
para("Approval to run a time-boxed proof of concept with access to representative "
     "source systems and a small client working group (a product owner, a "
     "knowledge/subject-matter contact per domain, and an IT/security contact). "
     "Deliverables: a working demonstration, a measured quality report, and a "
     "production reference architecture with a phased rollout plan.")

page_break()

# ===========================================================================
# 2. SOLUTION OVERVIEW
# ===========================================================================
h1("2.  Solution Overview")
para("Atlas turns a natural-language question into a single, grounded, cited "
     "answer in four steps. The user asks once; the platform does the finding, "
     "reading, and reconciling that people perform manually today.")
figure("fig3_query_flow.png",
       "Figure 1.  The \u201cask once\u201d journey: a single question becomes one "
       "consolidated, cited answer, with permissions enforced before any content "
       "reaches the model.")
olist([
    ("Unify.  ", "Connectors continuously ingest content from many systems into "
     "one normalized, searchable knowledge layer."),
    ("Understand.  ", "Content is parsed, chunked, embedded, and indexed for "
     "semantic and keyword retrieval, enriched with domain, project, owner, "
     "freshness, and access metadata."),
    ("Retrieve.  ", "A hybrid retriever and reranker find the most relevant "
     "passages across all sources, filtered to what the user may access."),
    ("Answer.  ", "A language model consolidates those passages into one grounded "
     "answer with inline citations, flagging conflicts and staleness, and "
     "declining when the evidence is insufficient."),
])
callout("Trust is a first-class feature: every answer is grounded in retrieved "
        "source text and cited, so it can be verified and audited — and Atlas says "
        "\u201cnot enough information\u201d rather than guessing.")

page_break()

# ===========================================================================
# 3. ARCHITECTURE
# ===========================================================================
h1("3.  Technical & Architectural Design")
para("This section describes the architecture that lets users ask a question once "
     "and receive a consolidated, source-cited answer synthesized across every "
     "relevant page, while remaining scalable, extensible, and secure.")

h2("3.1  Architecture goals and principles")
table(
    ["Goal", "Guiding principle"],
    [
        ["Centralize knowledge across domains and projects",
         "One normalized knowledge layer behind pluggable connectors."],
        ["Ask once, answer once",
         "A RAG pipeline that retrieves across all sources and synthesizes a single grounded answer with citations."],
        ["Scalable",
         "Stateless services, horizontally scalable workers, decoupled async pipelines, and a vector store built for billions of chunks."],
        ["Extensible",
         "A connector SDK, schema-driven metadata, swappable embedding/LLM models, and plug-in points for retrieval and answer formatting."],
        ["Trustworthy",
         "Every answer is grounded in retrieved text and cited; conflicts and staleness are surfaced."],
        ["Secure",
         "Permission-aware retrieval, data isolation, full auditability, and VPC/on-premises deployment."],
        ["Model-agnostic",
         "The LLM and embedding models are interfaces, not hard dependencies."],
    ],
    widths=[2.4, 4.1],
)
rich([("Non-goals (for the POC): ", True),
      ("replacing systems of record, editing source content, and fully automated "
       "write-back. Atlas is read-and-synthesize.", False)])

h2("3.2  High-level architecture")
para("Atlas is composed of four planes: an ingestion plane that turns source "
     "content into indexed knowledge; a knowledge plane of purpose-fit stores; a "
     "query plane that performs retrieval and answer synthesis (the RAG core); and "
     "an experience & platform plane providing APIs, UI, identity, security, and "
     "observability.")
figure("fig1_architecture.png",
       "Figure 2.  High-level architecture across the four planes. The query plane "
       "retrieves across all stores and synthesizes a single cited answer; identity "
       "and security wrap every interaction.")

h2("3.3  Ingestion plane")
h3("3.3.1  Connector framework")
para("Each source system is integrated through a connector that implements a "
     "small, well-defined contract. Connectors are the primary extensibility "
     "point — a new source is a new connector plus configuration, not a core "
     "change. A connector implements:")
bullet("enumerate the spaces, sites, projects, or folders it can see.", lead="discover() — ")
bullet("list documents and change markers since a checkpoint (incremental sync).",
       lead="list(since) — ")
bullet("return raw content and native metadata for one document.", lead="fetch(id) — ")
bullet("return the access-control descriptors used for permission-aware retrieval.",
       lead="acl(id) — ")
para("Connectors support scheduled incremental crawls, event-driven webhook "
     "updates, and one-time backfills. They are rate-limit aware, resumable via "
     "checkpoints, and run as independent, horizontally scalable workers.")

h3("3.3.2  Processing pipeline")
para("Each fetched document flows through an idempotent, queue-driven pipeline. "
     "Stages are independent workers, so they scale and fail independently.")
figure("fig2_pipeline.png",
       "Figure 3.  The ingestion and processing pipeline: each stage is an "
       "independent, idempotent worker.")
bullet("Text and structure extraction for HTML, PDF, Office formats, tables, and "
       "images (OCR); document structure is preserved to improve chunking and "
       "citation precision.")
bullet("Structure-aware chunking on semantic boundaries with controlled size and "
       "overlap, keeping a chunk \u2192 section \u2192 document hierarchy so "
       "citations point precisely.")
bullet("Deduplication and near-duplicate detection (content hashing plus embedding "
       "similarity) to collapse copies and flag conflicting versions.")
bullet("Metadata enrichment (domain, project, owner, freshness, language) and "
       "capture of the access-control descriptor at ingest time.")
bullet("Idempotent change detection: stable chunk IDs let re-processing update in "
       "place and remove orphaned chunks, handling edits and deletions at source.")

h2("3.4  Knowledge plane")
para("Atlas uses purpose-fit stores behind a single logical knowledge layer.")
table(
    ["Store", "Purpose", "Notes"],
    [
        ["Raw / object store", "Original bytes plus extracted text",
         "S3-compatible; cheap, durable, supports reprocessing."],
        ["Document & metadata store", "Canonical docs, chunks, metadata, ACLs, lineage",
         "PostgreSQL or document DB; source of truth for non-vector data."],
        ["Vector index", "Dense embeddings for semantic retrieval",
         "pgvector / OpenSearch / Milvus / Pinecone — pluggable."],
        ["Keyword index", "BM25 / lexical retrieval",
         "OpenSearch / Elasticsearch — for exact terms, codes, acronyms."],
        ["Knowledge graph (optional)", "Entities and relationships across documents",
         "Enables relationship questions and better disambiguation."],
        ["Cache", "Embeddings, hot queries, sessions", "Redis."],
    ],
    widths=[1.8, 2.4, 2.3],
)
callout("POC simplification: PostgreSQL with the pgvector extension can serve both "
        "metadata and vector roles, keeping the footprint small while preserving a "
        "clean path to dedicated stores at scale.")

h2("3.5  Query plane — the RAG core")
para("This is where \u201cask once, get a consolidated answer\u201d happens. The "
     "orchestrator interprets the question, retrieves across all sources using "
     "hybrid search, filters to the user\u2019s entitlements, reranks and "
     "diversifies the passages, and asks the language model to synthesize one "
     "grounded, cited answer.")
h3("3.5.1  Query understanding")
para("The orchestrator detects intent and scope (factual lookup, cross-project "
     "summary, comparison, or how-to), resolves filters such as domain, project, "
     "and time, and expands acronyms and synonyms — rewriting conversational "
     "follow-ups into standalone queries.")
h3("3.5.2  Hybrid retrieval")
para("Atlas combines dense (semantic) and sparse (keyword/BM25) retrieval and "
     "fuses the results. Semantic search captures meaning and paraphrase; keyword "
     "search nails exact terms — product codes, acronyms, IDs, policy numbers — "
     "that embeddings can miss. Retrieval is permission-filtered before anything "
     "reaches the model.")
h3("3.5.3  Reranking, deduplication, and diversification")
para("A cross-encoder reranker reorders candidates by true relevance; near-"
     "identical passages are deduplicated and results are diversified across "
     "documents and sources, so the consolidated answer reflects all relevant "
     "pages rather than several copies of the same one. A token budget bounds the "
     "context.")
h3("3.5.4  Answer synthesis with citations")
para("The model receives the question and the curated passages (each tagged with a "
     "citation ID) under a strict instruction to answer only from the provided "
     "passages, cite each claim, consolidate across sources, note conflicts and "
     "staleness, and respond \u201cnot enough information\u201d when the context "
     "does not support an answer. Output is structured: the answer (with inline "
     "citation markers), a list of citations (title, URL, snippet, freshness, "
     "score), and a confidence/coverage signal. Streaming keeps the experience "
     "responsive.")
callout("Anti-hallucination posture: grounding + citations + \u201cinsufficient "
        "evidence\u201d handling + post-hoc citation verification + a continuous "
        "evaluation harness. Trust is engineered, not assumed.")

h2("3.6  Scalability")
para("Ingestion (throughput-bound) and query (latency-bound) scale independently "
     "and never block one another.")
figure("fig5_scalability.png",
       "Figure 4.  Decoupled planes scale independently: add workers for ingestion "
       "throughput; add stateless query instances for concurrency.")
bullet("Async, queue-driven ingestion: each stage is an idempotent worker; "
       "back-pressure and retries are handled by the queue. Scale by adding workers.")
bullet("Stateless query services scale horizontally behind a load balancer; "
       "session state lives in Redis or the database.")
bullet("Vector store scaling via approximate-nearest-neighbour indexes (HNSW/IVF) "
       "with sharding and replication; the store is an interface, so moving from "
       "pgvector (POC) to Milvus/OpenSearch (scale) requires no application change.")
bullet("Caching of embeddings, query results, and reranks; tiered models and token "
       "budgeting control cost and latency.")
para("Target POC scale is hundreds of thousands to low millions of chunks; the "
     "same design extends to hundreds of millions and beyond by swapping in "
     "dedicated stores and adding workers — with no architectural change.")

h2("3.7  Security, permissions, and governance")
para("Security is central because Atlas reads sensitive enterprise knowledge.")
figure("fig4_permissions.png",
       "Figure 5.  Permission-aware retrieval: candidate documents are intersected "
       "with the user\u2019s entitlements before any content reaches the model.",
       width=5.6)
bullet("Authentication via SSO (OIDC/SAML); the platform never holds primary "
       "credentials.", lead="Authentication.  ")
bullet("Each document\u2019s ACL is captured at ingest, and at query time results "
       "are filtered to the intersection of the user\u2019s roles and the document "
       "ACLs, so the model only ever receives content the user may read.",
       lead="Permission-aware retrieval.  ")
bullet("ACLs refresh on schedule and on change events; sensitive sources can be "
       "re-validated at query time to avoid stale-grant leaks.",
       lead="Permission freshness.  ")
bullet("Encryption in transit and at rest, secrets in a vault/KMS, and optional "
       "PII detection and redaction in the pipeline.", lead="Data protection.  ")
bullet("Runs in the client\u2019s VPC, private cloud, or on-premises; models can "
       "be self-hosted so no content leaves the boundary.", lead="Deployment flexibility.  ")
bullet("Every query, retrieval set, and answer is logged for compliance and "
       "debugging, with access controls on the logs themselves.", lead="Auditability.  ")

h2("3.8  Extensibility")
para("Growth is configuration, not re-architecture. Everything the query plane "
     "consumes is an interface — vector store, keyword store, embedding model, "
     "reranker, and LLM are all swappable.")
table(
    ["Extension point", "How it extends", "Example"],
    [
        ["Connectors", "Implement the connector contract", "Add ServiceNow, GitHub, or a custom REST API."],
        ["Domains / projects", "Schema-driven metadata and tagging", "Onboard a new business domain with its taxonomy."],
        ["Embedding model", "Embedding interface", "Swap to a domain-tuned or newer model; re-embed."],
        ["LLM", "Synthesizer interface", "Switch providers, self-host, or route by query type."],
        ["Retrieval strategy", "Pluggable retrievers/rerankers", "Add graph retrieval or multi-hop reasoning."],
        ["Answer formats", "Output templates", "Brief, detailed-with-citations, comparison table, checklist."],
        ["Surfaces", "API-first", "Web app, chat, Slack/Teams, IDE, embedded widgets."],
    ],
    widths=[1.6, 2.3, 2.6],
)

h2("3.9  Observability, evaluation, and quality")
para("Trustworthy answers must be measured. Atlas ships with end-to-end tracing "
     "and a continuous RAG evaluation harness.")
bullet("OpenTelemetry traces across ingestion and query (which chunks were "
       "retrieved, scores, tokens, latency per stage).", lead="Tracing.  ")
bullet("Ingestion throughput and freshness, query latency (p50/p95), retrieval hit "
       "rates, cache rates, and model cost.", lead="Metrics.  ")
bullet("Retrieval quality (recall@k, MRR, nDCG), answer quality (groundedness, "
       "relevance, citation correctness), and safety (hallucination rate, correct "
       "refusals) on a curated benchmark.", lead="Evaluation.  ")
bullet("Thumbs up/down and comments feed evaluation and prioritize fixes; "
       "model/prompt/retrieval changes are gated against the suite before "
       "promotion.", lead="Feedback & gating.  ")

h2("3.10  Reference technology stack")
para("The stack is a recommendation; each layer is swappable and final choices "
     "adapt to the client\u2019s existing platform and constraints.")
table(
    ["Layer", "POC recommendation", "Scale / alternatives"],
    [
        ["Connectors & pipeline", "Python workers (FastAPI control plane)", "Same, scaled horizontally"],
        ["RAG orchestration", "LangChain / LlamaIndex or thin custom orchestrator", "Custom orchestrator"],
        ["Queue / async", "Redis Streams / RabbitMQ", "Kafka / managed queues"],
        ["Metadata store", "PostgreSQL", "PostgreSQL (HA) / managed"],
        ["Vector index", "PostgreSQL + pgvector", "Milvus / OpenSearch / Pinecone"],
        ["Keyword index", "OpenSearch / Elasticsearch", "Same, clustered"],
        ["Object store", "S3-compatible (MinIO/S3)", "Cloud object storage"],
        ["Embeddings & LLM", "Managed or self-hosted", "Self-hosted for data residency"],
        ["Identity", "OIDC/SAML via existing IdP", "Same"],
        ["Frontend", "React / Next.js streaming chat", "Same"],
        ["Deployment", "Docker + Kubernetes (Helm)", "K8s autoscaling, multi-AZ"],
        ["Observability", "OpenTelemetry + Prometheus/Grafana", "+ tracing/eval tooling"],
    ],
    widths=[1.7, 2.5, 2.3],
)

h2("3.11  Deployment view")
para("All components run inside the client\u2019s security boundary. The only "
     "optional egress is to a managed model API, which can be removed entirely by "
     "self-hosting the models.")
figure("fig6_deployment.png",
       "Figure 6.  Deployment inside the client VPC or on-premises, with stateful "
       "services and optional self-hosted model serving.")

h2("3.12  Key non-functional targets (POC)")
table(
    ["Attribute", "Target (POC)"],
    [
        ["Answer latency", "First token in a few seconds; full answer typically under ~10s."],
        ["Retrieval recall@10", "High on the curated benchmark (tuned during the POC)."],
        ["Citation correctness", "Majority of claims correctly supported; measured and improved."],
        ["Freshness", "Source changes reflected within minutes (events) to hours (crawl)."],
        ["Availability", "Single-region HA for the POC; multi-AZ path for production."],
        ["Security", "100% permission-aware retrieval; no cross-entitlement leakage."],
    ],
    widths=[1.9, 4.6],
)

page_break()

# ===========================================================================
# 4. POC PLAN
# ===========================================================================
h1("4.  Proof-of-Concept Plan")

h2("4.1  Objectives")
para("Prove, end to end and on a representative slice of real content, that Atlas "
     "can:")
olist([
    "unify heterogeneous knowledge sources behind one knowledge layer;",
    "answer cross-document questions with a single, consolidated, cited answer;",
    "enforce permission-aware retrieval with no cross-entitlement leakage;",
    "meet measurable quality and latency targets on a curated benchmark; and",
    "demonstrate a credible path to scale and extend to all sources and domains.",
])

h2("4.2  Scope")
h3("In scope")
bullet("Two to three connectors to representative sources (for example a wiki, a "
       "file share/SharePoint, and one structured source such as a ticketing system).")
bullet("One to two priority domains/projects chosen with the client for high value "
       "and manageable size.")
bullet("The full ingestion pipeline and hybrid retrieval, reranking, and grounded "
       "answer synthesis with citations.")
bullet("Permission-aware retrieval honoring source ACLs for a defined set of test "
       "users and groups.")
bullet("A polished web chat/search UI with expandable source citations.")
bullet("An evaluation harness, a curated benchmark question set, and a results "
       "report; plus a reference architecture and rollout plan.")
h3("Out of scope")
bullet("Connecting every source and domain (proven extensible; rolled out after the POC).")
bullet("Editing or writing back to source systems.")
bullet("Fully autonomous agents and workflows (a clear future extension).")
bullet("Production-grade HA/DR hardening (designed for, not fully built, in the POC).")

h2("4.3  Phased plan")
para("Phases are sequenced by dependency, not by calendar; each has explicit "
     "entry/exit criteria so progress is measured by completed capability.")
table(
    ["Phase", "Focus", "Key outputs"],
    [
        ["0 — Discovery & setup", "Confirm sources, domains, test users, and metrics; provision the environment in the client VPC; security review.", "Signed-off scope, environment, access."],
        ["1 — Ingestion & indexing", "Build/configure connectors; run backfill and incremental sync; populate stores.", "Indexed corpus, ingestion metrics, freshness."],
        ["2 — Retrieval & answers", "Hybrid retrieval, reranking, grounded synthesis with citations; UI.", "Working Q&A demo over the corpus."],
        ["3 — Security & permissions", "ACL capture and permission-aware retrieval; isolation; audit logging.", "Permission test suite passing."],
        ["4 — Evaluation & tuning", "Build the benchmark; measure; tune chunking, retrieval, and prompts.", "Quality report versus targets."],
        ["5 — Demo & decision", "Stakeholder demo; results review; production plan and costing.", "Go/no-go package."],
    ],
    widths=[1.6, 3.1, 1.8],
)

h2("4.4  Success criteria")
para("The POC is successful if, on the curated benchmark and live demonstration, "
     "the following hold. Exact numeric thresholds are agreed in Phase 0.")
table(
    ["#", "Criterion", "How measured"],
    [
        ["1", "Cross-document consolidated answers", "Answers synthesize from multiple sources, not single-document lookups."],
        ["2", "Grounded and cited", "Every substantive claim carries a citation; citation correctness above threshold."],
        ["3", "Retrieval quality", "recall@k / nDCG on the benchmark above threshold."],
        ["4", "Permission safety", "100% of permission tests pass; no test user receives inaccessible content."],
        ["5", "Latency", "First token within a few seconds; full answers within the agreed bound."],
        ["6", "Freshness", "Source edits/deletes reflected within the agreed window."],
        ["7", "Extensibility shown", "Adding a new connector/domain demonstrated as configuration plus a small connector."],
        ["8", "Stakeholder validation", "SMEs rate answers useful/accurate on a sample, with a positive majority."],
    ],
    widths=[0.4, 2.1, 4.0],
)

h2("4.5  Evaluation methodology")
bullet("A benchmark of 50–150 real questions per domain, authored with SMEs, each "
       "with an expected answer and expected source documents.")
bullet("Automated metrics for retrieval (recall@k, MRR, nDCG) and answers "
       "(groundedness, relevance, citation correctness, hallucination rate, correct refusals).")
bullet("Human review of a sample by SMEs, plus in-product thumbs up/down feeding the loop.")
bullet("Regression gating: changes to chunking, retrieval, prompts, or models are "
       "re-scored before adoption, with results on a single quality dashboard.")

h2("4.6  What we need from the client")
table(
    ["Need", "Why"],
    [
        ["Read access to 2–3 representative sources (service accounts)", "Ingestion."],
        ["Choice of 1–2 priority domains/projects", "Focus and value."],
        ["3–5 test users/groups with varied permissions", "Permission validation."],
        ["SME time for the benchmark and answer review", "Quality measurement."],
        ["Security/IT contact and an environment in the client VPC", "Secure deployment and review."],
        ["A product owner / sponsor", "Decisions and prioritization."],
    ],
    widths=[3.7, 2.8],
)

h2("4.7  Team model (lean POC)")
table(
    ["Role", "Responsibility"],
    [
        ["Product owner (client)", "Priorities, scope decisions, stakeholder access."],
        ["Solution architect", "Architecture, integration, and security alignment."],
        ["ML/RAG engineer", "Retrieval, embeddings, synthesis, and evaluation."],
        ["Backend engineer", "Connectors, pipeline, APIs, and stores."],
        ["Frontend engineer (part-time)", "Chat/search UI."],
        ["Subject-matter expert(s) (client)", "Domain validation and benchmark authoring."],
        ["Security / IT (client)", "Access, deployment, and compliance review."],
    ],
    widths=[2.4, 4.1],
)

h2("4.8  Cost considerations")
para("POC cost is modest and dominated by a small Kubernetes footprint plus "
     "stateful stores (which can run in the client\u2019s existing cloud/VPC); "
     "model usage for embeddings and per-query synthesis, controlled via tiered "
     "models, caching, batching, and token budgeting; and the lean team above. "
     "Self-hosting models trades API cost for compute and removes egress concerns. "
     "A per-query and monthly estimate is provided after Phase 1, once corpus size "
     "and query volume are validated.")

h2("4.9  Risks and mitigations")
table(
    ["Risk", "Mitigation"],
    [
        ["Source access / API limits delay ingestion", "Start access requests in Phase 0; rate-limit-aware, resumable connectors."],
        ["Answer quality / hallucination below bar", "Grounding, citations, refusal handling, citation verification, and evaluation gating; tune chunking and retrieval."],
        ["Permission leakage (severe)", "Permission-aware retrieval as a core control; dedicated permission test suite; late-binding checks; audit logs."],
        ["Messy or inconsistent content", "Deduplication, conflict flagging, and freshness signals; surface — not hide — disagreements."],
        ["Data residency / privacy", "VPC/on-prem deployment; self-hostable models; PII detection/redaction; encryption and KMS."],
        ["Scope creep (connect everything now)", "Strict POC scope; extensibility proven, full rollout phased afterwards."],
        ["Cost surprises at scale", "Tiered models, caching, batching; cost estimate after Phase 1; cost dashboard."],
        ["Vendor / model lock-in", "Model-agnostic interfaces; swappable stores, embeddings, and LLM."],
        ["User adoption", "Polished UX; citations build trust; feedback loop; change management in rollout."],
    ],
    widths=[2.5, 4.0],
)

h2("4.10  From POC to production")
para("Once the POC clears the go/no-go decision, the path to a durable, "
     "organization-wide knowledge asset is:")
olist([
    "broaden connectors and domains (configuration plus connectors), with "
    "per-source freshness and ACL policies;",
    "harden for scale — dedicated vector/keyword stores, more workers, "
    "autoscaling, multi-AZ, HA/DR, and backups;",
    "operationalize quality — continuous evaluation, regression gating, and "
    "feedback-driven improvement;",
    "expand surfaces — chat, Slack/Teams, IDE, and embedded widgets via the "
    "API-first design; and",
    "extend capabilities — comparison/table answers, multi-hop reasoning, "
    "scheduled briefings, and carefully governed agentic workflows.",
])

page_break()

# ===========================================================================
# 5. CONCLUSION
# ===========================================================================
h1("5.  Conclusion & Recommendation")
para("Knowledge across the organization is valuable but fragmented, and the cost "
     "is paid every day in the manual effort of finding, reading, and reconciling "
     "documents. Atlas addresses this directly: a scalable, extensible platform "
     "that centralizes knowledge across domains and projects and lets people ask a "
     "question once and receive one clear, consolidated, cited answer — instead of "
     "reading documents one by one.")
para("The architecture is deliberately pragmatic: decoupled planes that scale "
     "independently, swappable components that avoid vendor lock-in, "
     "permission-aware retrieval that protects sensitive content, and a measured "
     "evaluation harness that makes answer quality a managed, observable property "
     "rather than a hope.")
rich([("Recommendation: ", True),
      ("proceed with the time-boxed proof of concept defined in Section 4. It is "
       "low-risk and high-signal — delivering a working demonstration, a measured "
       "quality report, and a production-ready reference architecture, and giving "
       "the client an evidence-based decision with a clear, phased path to full "
       "rollout.", False)])
callout("Ask once. Receive one clear, consolidated, trustworthy answer — across "
        "every domain and project.")

# ===========================================================================
# APPENDIX
# ===========================================================================
h1("Appendix A — Glossary")
table(
    ["Term", "Meaning"],
    [
        ["RAG (retrieval-augmented generation)", "Retrieving relevant source passages and using a language model to synthesize a grounded answer from them."],
        ["Embedding", "A numeric vector representing the meaning of text, enabling semantic similarity search."],
        ["Vector index", "A store and search index for embeddings using approximate-nearest-neighbour algorithms."],
        ["Chunk", "A bounded, structure-aware passage of a document used as the unit of retrieval and citation."],
        ["Hybrid retrieval", "Combining semantic (vector) and keyword (BM25) search and fusing the results."],
        ["Reranker", "A model that reorders retrieved candidates by their true relevance to the question."],
        ["ACL (access-control list)", "The descriptor of who may read a document, captured at ingest and enforced at query time."],
        ["Groundedness", "The degree to which an answer is supported by the retrieved source passages."],
        ["Connector", "A pluggable component that ingests content and permissions from one source system."],
    ],
    widths=[2.2, 4.3],
)

# ---------------------------------------------------------------------------
# Apply header/footer to all sections
# ---------------------------------------------------------------------------
for section in doc.sections:
    add_footer(section)
    add_header(section, "Atlas — Unified Knowledge Platform")
# Remove header/footer from the cover page via a different first-page setting
doc.sections[0].different_first_page_header_footer = True
fp_footer = doc.sections[0].first_page_footer
fp_footer.paragraphs[0].text = ""
fp_header = doc.sections[0].first_page_header
fp_header.paragraphs[0].text = ""

# Ask Word to refresh fields (TOC, PAGE, NUMPAGES) when the document is opened.
settings = doc.settings.element
upd = OxmlElement("w:updateFields")
upd.set(qn("w:val"), "true")
settings.append(upd)

doc.save(OUT)
print("wrote", os.path.normpath(OUT))
